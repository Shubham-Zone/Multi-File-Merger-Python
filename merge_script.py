import os
from docx import Document

def handle_file(doc, file_path):
    doc.add_heading(file_path, level=1)  # Add file name as heading
    with open(file_path, 'r', encoding='utf-8') as file:
        content = file.read()
        doc.add_paragraph(content)

def merge_files_to_docx(folder_path, file_extension, output_file):
    doc = Document()

    # Iterate over files in the folder
    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
        if filename.endswith(file_extension):
            handle_file(doc, file_path)

    # Save the merged content to a Word document
    doc.save(output_file)

# Example usage
folder_path = r'E:\STUDY\CODING\MCA\SEM 2\ARTIFICIAL INTELLIGENCE\PR\PR_FILE'
file_extension = '.pl'  # Change this to the desired file extension
output_file = 'merged_document.docx'
merge_files_to_docx(folder_path, file_extension, output_file)
